"""
DocxExporter — 报告导出为 Word (.docx) 文件
=========================================
将报告内容（Markdown 文本）+ chart 图片嵌入为 Word 文档。
使用 Word 内置 Heading 样式（支持自动生成目录）。

遵循 Hermes Code Rules 规范
"""
from __future__ import annotations

import logging
import re
import struct
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import TYPE_CHECKING, Any

if TYPE_CHECKING:
    from docx.text.paragraph import Paragraph

import hashlib

import requests as _requests

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement
from docx.shared import Inches, Pt, RGBColor
from docx.image.exceptions import UnrecognizedImageError

logger = logging.getLogger(__name__)


def _safe_unlink(path: Path) -> None:
    """删除文件；失败时仅记 warning 并忽略。

    在删除受限的场景（网络盘、只读文件、安全软件/回收站拦截）下优雅降级，
    避免删除失败导致整个生成/清理流程抛异常。
    """
    try:
        path.unlink()
    except OSError as e:
        logger.warning("无法删除文件（已忽略）: %s: %s", path, e)


# ── 图片尺寸自适应 ────────────────────────────────────────

# A4 横向内容区约 6.5 英寸；纵向约 4.8 英寸。统一上限：
_MAX_IMG_WIDTH_INCH = 5.5
_MAX_IMG_HEIGHT_INCH = 7.0


def _read_image_size(img_path: Path) -> tuple[int, int] | None:
    """读取图片宽高（像素）。支持 PNG / JPEG / GIF，不依赖 Pillow。

    失败返回 None，调用方按固定宽度处理。
    """
    try:
        with open(img_path, "rb") as f:
            head = f.read(32)
        if head.startswith(b"\x89PNG\r\n\x1a\n"):
            # PNG: IHDR 起始在 16 字节处
            with open(img_path, "rb") as f:
                f.read(16)  # skip signature + IHDR chunk header
                w, h = struct.unpack(">II", f.read(8))
            return (w, h)
        if head.startswith(b"\xff\xd8"):
            # JPEG: 扫描 SOF0 (0xFFC0) / SOF2 (0xFFC2) 段
            with open(img_path, "rb") as f:
                f.read(2)  # skip SOI marker
                # 限制最大扫描段数，防止恶意文件导致死循环
                max_segments = 64
                for _ in range(max_segments):
                    marker = f.read(2)
                    if len(marker) < 2:
                        return None
                    # 跳过填充字节 0xFF
                    while marker == b"\xff\xff":
                        marker = marker[1:2] + f.read(1)
                        if len(marker) < 2:
                            return None
                    if marker[0] != 0xFF:
                        # 非 marker，可能文件损坏
                        return None
                    code = marker[1]
                    # SOF0/SOF2: 找到帧头
                    if code in (0xC0, 0xC2):
                        f.read(3)  # skip length + precision
                        h, w = struct.unpack(">HH", f.read(4))
                        return (w, h)
                    # SOS（Start of Scan）：后面是压缩数据，没 SOF 了
                    if code == 0xDA:
                        return None
                    # RSTn (0xD0-0xD7) 和 SOI (0xD8) / EOI (0xD9) 没有段长度
                    # 直接跳过 marker 本身即可
                    if 0xD0 <= code <= 0xD9:
                        continue
                    # 其他标记：读取段长度并跳过
                    seg_len_bytes = f.read(2)
                    if len(seg_len_bytes) < 2:
                        return None
                    seg_len = struct.unpack(">H", seg_len_bytes)[0]
                    # seg_len 包含自身 2 字节，至少为 2；< 2 表示文件损坏
                    if seg_len < 2:
                        return None
                    f.seek(seg_len - 2, 1)
                return None  # 超过最大段数仍未找到 SOF
        if head[:6] in (b"GIF87a", b"GIF89a"):
            with open(img_path, "rb") as f:
                f.read(6)
                w, h = struct.unpack("<HH", f.read(4))
            return (w, h)
    except (OSError, struct.error) as e:
        logger.debug("read image size failed for %s: %s", img_path, e)
    return None


def _fit_image_size(img_path: Path) -> tuple[Any, Any]:
    """根据图片实际宽高比返回 (width, height)。

    竖图（高>宽）按高度限制，避免压扁；横图按宽度限制。
    python-docx 可只给 width 或 height，另一个自动按比例。这里返回其中一个。
    """
    size = _read_image_size(img_path)
    if size is None:
        return (Inches(_MAX_IMG_WIDTH_INCH), None)

    w_px, h_px = size
    aspect = w_px / h_px if h_px > 0 else 1.0

    if aspect >= 1.0:
        # 横图或方图：按宽度限制
        return (Inches(_MAX_IMG_WIDTH_INCH), None)
    # 竖图：按高度限制
    return (None, Inches(_MAX_IMG_HEIGHT_INCH))


_MAX_DOWNLOAD_BYTES = 20 * 1024 * 1024  # 20MB，防止恶意大文件
_MAX_REDIRECTS = 5  # 重定向上限


def _download_image(url: str, dest_dir: Path) -> Path | None:
    """下载 HTTP/HTTPS 图片到 dest_dir，返回本地路径。

    失败返回 None。

    安全措施：
    - 单次 socket 操作超时 30s
    - 重定向上限 5 次（防止重定向环）
    - 响应体上限 20MB（防止恶意大文件）
    """
    # SSRF 加固：仅允许 http/https 协议
    if not url.lower().startswith(("http://", "https://")):
        logger.warning("不支持的图片 URL 协议（仅允许 http/https）: %s", url[:80])
        return None

    # 从 URL 推断扩展名
    url_path = url.split("?")[0].split("/")[-1]
    ext = Path(url_path).suffix or ".png"
    if ext.lower() not in (".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp"):
        ext = ".png"

    # 用 url 的 md5 作为文件名，避免重复下载
    url_hash = hashlib.md5(url.encode("utf-8")).hexdigest()[:12]
    dest_path = dest_dir / f"downloaded_{url_hash}{ext}"

    if dest_path.exists():
        # 验证缓存文件完整性：太小可能是上次下载失败的残留
        if dest_path.stat().st_size < 32:
            logger.warning("缓存文件过小 (%d bytes)，可能损坏，重新下载: %s",
                           dest_path.stat().st_size, dest_path.name)
            _safe_unlink(dest_path)
        else:
            return dest_path

    try:
        dest_dir.mkdir(parents=True, exist_ok=True)
        # 用 with 包裹 response 确保连接释放（stream=True 下尤其重要）
        with _requests.get(url, timeout=30, stream=True,
                           headers={"User-Agent": "Mozilla/5.0"}) as resp:
            # 检查重定向次数：history 长度 > _MAX_REDIRECTS 才算超限
            # （history 记录的是中间跳转，5 次跳转后到达终点的仍允许）
            if len(resp.history) > _MAX_REDIRECTS:
                logger.warning("重定向次数过多 (%d)，放弃: %s",
                               len(resp.history), url[:80])
                return None
            resp.raise_for_status()
            # SSRF 加固：若服务器明确返回非图片类型，拒绝落盘
            content_type = getattr(resp, "headers", {}).get("Content-Type", "")
            if content_type and not content_type.lower().startswith("image/"):
                logger.warning("响应非图片类型 (%s)，放弃: %s", content_type, url[:80])
                return None
            # 流式下载，限制最大字节数
            data = b""
            for chunk in resp.iter_content(chunk_size=65536):
                if chunk:
                    data += chunk
                    if len(data) > _MAX_DOWNLOAD_BYTES:
                        logger.warning("图片过大，已超过 %d 字节，放弃: %s",
                                       _MAX_DOWNLOAD_BYTES, url[:80])
                        return None
            dest_path.write_bytes(data)
        logger.info("下载图片: %s → %s", url[:80], dest_path.name)
        return dest_path
    except (_requests.RequestException, OSError) as e:
        logger.warning("下载图片失败 %s: %s", url[:80], e)
        return None


# ── 目录字段 XML（Word TOC） ────────────────────────────

_TOC_XML = (
    '<w:sdt xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
    '<w:sdtPr>'
    '<w:docPartObj><w:docPartGallery w:val="Table of Contents"/></w:docPartObj>'
    '</w:sdtPr>'
    '<w:sdtContent>'
    '<w:p><w:pPr><w:jc w:val="center"/></w:pPr>'
    '<w:r><w:t>（打开 Word 后在目录上右键 → 更新域 → 更新整个目录）</w:t></w:r>'
    '</w:p>'
    '<w:p><w:r><w:br/></w:r></w:p>'
    '</w:sdtContent></w:sdt>'
)


def _add_heading(doc: Document, text: str, level: int) -> None:
    """使用 Word 内置 Heading 样式添加标题。"""
    level = min(level, 9)
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def _add_body(doc: Document, text: str) -> None:
    """添加正文段落（支持 **bold** 内联标记）。"""
    if not text.strip():
        return
    p = doc.add_paragraph()
    _add_formatted_run(p, text.strip())
    p.paragraph_format.first_line_indent = Pt(22)  # 首行缩进约两个中文字符（11pt * 2）
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = Pt(18)


def _add_blockquote(doc: Document, text: str) -> None:
    """添加引用块（灰色背景、缩进）。"""
    if not text.strip():
        return
    p = doc.add_paragraph()
    _add_formatted_run(p, text.strip())
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    # 添加左侧灰色边框（模拟引用样式）
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        '<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:left w:val="single" w:sz="12" w:space="8" w:color="999999"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)


def _add_inline_code_run(p: "Paragraph", code: str) -> None:
    """添加行内代码 run：Consolas 字体 + 浅灰背景。"""
    run = p.add_run(code)
    run.font.size = Pt(10)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    # 浅灰背景（通过 w:shd 实现）
    rPr = run._element.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F0F0F0")
    rPr.append(shd)


def _add_hyperlink_run(p: "Paragraph", text: str, url: str) -> None:
    """添加超链接 run（蓝色 + 下划线 + 可点击）。"""
    part = p.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    # 蓝色
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)
    # 下划线
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Microsoft YaHei")
    rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    rPr.append(rFonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "22")  # 11pt = 22 half-points
    rPr.append(sz)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    p._element.append(hyperlink)


def _add_formatted_run(p: "Paragraph", text: str) -> None:
    """向段落添加文本，解析行内标记：

    - **bold** / __bold__ → 加粗 run
    - *italic* / _italic_ → 斜体 run
    - `code` → Consolas + 浅灰背景 run
    - ~~text~~ → 删除线 run
    - [text](url) → 超链接 run
    - <br> → 换行

    不支持嵌套。`_italic_` / `__bold__` 形式要求两侧不在单词中间，
    以避免 snake_case_variable 等被误判。
    """
    # 用统一正则切分所有标记类型
    # `_italic_` / `__bold__` 加词边界断言，防止变量名被误判
    # link 部分 `(?:[^()]|\([^()]*\))*` 允许 URL 中包含一层嵌套括号
    # （例如 Wikipedia: https://en.wikipedia.org/wiki/C_(programming_language)）
    pattern = re.compile(
        r"(?P<bold>\*\*[^*]+\*\*|(?<!\w)__(?!_)[^_]+(?<!_)__(?!\w))"
        r"|(?P<italic>\*[^*]+\*|(?<!\w)_(?!_)[^_]+(?<!_)_(?!\w))"
        r"|(?P<code>`[^`]+`)"
        r"|(?P<strike>~~[^~]+~~)"
        r"|(?P<link>\[[^\]]+\]\((?:[^()]|\([^()]*\))*\))"
        r"|(?P<br><br\s*/?>)"
    )

    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            # 普通文本
            plain = text[pos:m.start()]
            if plain:
                _add_plain_run(p, plain)

        if m.group("bold"):
            raw = m.group("bold")
            # **bold** 取 [2:-2]，__bold__ 取 [2:-2]
            inner = raw[2:-2]
            run = p.add_run(inner)
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.name = "Microsoft YaHei"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        elif m.group("italic"):
            raw = m.group("italic")
            # *italic* / _italic_ 均取 [1:-1]
            inner = raw[1:-1]
            run = p.add_run(inner)
            run.font.size = Pt(11)
            run.font.italic = True
            run.font.name = "Microsoft YaHei"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        elif m.group("code"):
            inner = m.group("code")[1:-1]
            _add_inline_code_run(p, inner)
        elif m.group("strike"):
            inner = m.group("strike")[2:-2]
            run = p.add_run(inner)
            run.font.size = Pt(11)
            run.font.strike = True
            run.font.name = "Microsoft YaHei"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        elif m.group("link"):
            link_text = m.group("link")
            # [text](url)，URL 可能包含一层嵌套括号
            m_link = re.match(
                r"\[([^\]]+)\]\(((?:[^()]|\([^()]*\))*)\)",
                link_text,
            )
            if m_link:
                _add_hyperlink_run(p, m_link.group(1), m_link.group(2))
        elif m.group("br"):
            run = p.add_run()
            run.add_break()
        pos = m.end()

    # 末尾普通文本
    if pos < len(text):
        plain = text[pos:]
        if plain:
            _add_plain_run(p, plain)


def _add_plain_run(p: "Paragraph", text: str) -> None:
    """添加普通文本 run。"""
    if not text:
        return
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def _add_bullet(doc: Document, text: str, *, level: int = 0) -> None:
    """添加无序列表项（支持 **bold** 等内联标记 + 嵌套层级）。"""
    p = doc.add_paragraph(style="List Bullet")
    _add_formatted_run(p, text)
    if level > 0:
        p.paragraph_format.left_indent = Inches(0.25 * level)


def _add_numbered(doc: Document, text: str, *, level: int = 0) -> None:
    """添加有序列表项（支持嵌套层级）。"""
    p = doc.add_paragraph(style="List Number")
    _add_formatted_run(p, text)
    if level > 0:
        p.paragraph_format.left_indent = Inches(0.25 * level)


def _add_task_list_item(doc: Document, text: str, *, checked: bool = False, level: int = 0) -> None:
    """添加任务列表项：☐ 未完成 / ☑ 已完成。

    使用 Unicode 符号 + 普通段落样式（非 List Bullet），以保持复选框可见。
    """
    p = doc.add_paragraph()
    checkbox = "☑" if checked else "☐"
    cb_run = p.add_run(f"{checkbox} ")
    cb_run.font.size = Pt(12)
    cb_run.font.name = "Microsoft YaHei"
    cb_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    _add_formatted_run(p, text)
    if level > 0:
        p.paragraph_format.left_indent = Inches(0.25 * level)


def _parse_code_block(doc: Any, lines: list[str], idx: int) -> int:
    """解析围栏代码块（```language ... ```），返回消耗的行数。

    - 普通代码块：灰色背景 + Consolas 等宽字体
    - mermaid 代码块：提示文字说明（Word 原生不支持 Mermaid 渲染）
    """
    first_line = lines[idx].strip()
    lang = first_line[3:].strip().lower() if len(first_line) > 3 else ""

    code_lines: list[str] = []
    idx += 1
    while idx < len(lines):
        line = lines[idx]
        stripped = line.strip()
        if stripped.startswith("```"):
            idx += 1
            break
        code_lines.append(line)
        idx += 1

    code_text = "\n".join(code_lines)

    # Mermaid 代码块：用提示框展示（Word 不能直接渲染 Mermaid）
    if lang == "mermaid":
        _add_blockquote(doc, f"[Mermaid 图表 · {len(code_lines)} 行]（Word 不支持直接渲染，需手动转图后插入）")
        return idx

    # 普通代码块：灰色背景 + 等宽字体
    # 使用一个单元格的表格来实现背景色
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.rows[0].cells[0]
    cell.width = Inches(6.5)

    # 设置单元格底纹（浅灰）
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F5F5F5")
    tc_pr.append(shd)

    # 清空默认段落，添加代码内容
    cell.paragraphs[0].clear()
    if lang:
        run = cell.paragraphs[0].add_run(f"# {lang}")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        cell.add_paragraph("")

    for j, code_line in enumerate(code_lines):
        # 空行保留为真空行（不再用空格替代）
        if j == 0 and not lang:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        if code_line:  # 非空行才添加 run，空行保持段落为空
            run = p.add_run(code_line)
            run.font.size = Pt(9)
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")

    # 代码块后空行
    doc.add_paragraph("")

    return idx


def _parse_alignment_row(stripped: str) -> list[WD_ALIGN_PARAGRAPH] | None:
    """解析 markdown 表格分隔行，返回各列对齐方式列表。

    `|:---|` → LEFT, `|:---:|` → CENTER, `|---:|` → RIGHT
    支持空格: `| :--- | ---: |`
    允许部分单元格为空（默认 LEFT 对齐）。
    若不是分隔行返回 None。
    """
    if not stripped.startswith("|"):
        return None
    # 必须包含至少一个 dash 或 colon 才可能是分隔行
    if "-" not in stripped and ":" not in stripped:
        return None
    # 拆分单元格（去除首尾的 | 之后按 | 分隔）
    cells = [c.strip() for c in stripped.strip().strip("|").split("|")]
    aligns: list[WD_ALIGN_PARAGRAPH] = []
    has_separator = False  # 至少有一个 cell 看起来像分隔符
    for c in cells:
        if not c:
            # 空单元格：默认左对齐
            aligns.append(WD_ALIGN_PARAGRAPH.LEFT)
            continue
        # 合法分隔符格式：:?[-:]+:? （至少 1 个 dash/colon，可选首尾 colon）
        if not re.match(r"^:?[-:]+:?$", c):
            # 任一单元格不是合法分隔符 → 整行不是分隔行
            return None
        has_separator = True
        if c.startswith(":") and c.endswith(":"):
            aligns.append(WD_ALIGN_PARAGRAPH.CENTER)
        elif c.endswith(":"):
            aligns.append(WD_ALIGN_PARAGRAPH.RIGHT)
        else:
            aligns.append(WD_ALIGN_PARAGRAPH.LEFT)
    if not has_separator:
        return None
    return aligns


def _parse_markdown_table(doc: Any, line: str, lines: list[str], idx: int) -> int:
    """解析 markdown 表格，返回消耗的行数。支持列对齐方式。"""
    idx_init = idx  # 记录起始位置，用于回退渲染
    rows_text: list[list[str]] = []
    col_aligns: list[WD_ALIGN_PARAGRAPH] | None = None

    while idx < len(lines):
        stripped = lines[idx].strip()
        if not stripped.startswith("|"):
            break
        # 分隔行：解析对齐方式
        aligns = _parse_alignment_row(stripped)
        if aligns is not None:
            col_aligns = aligns
            idx += 1
            continue
        # 数据行：允许有无尾部 |（GFM 规范尾部 pipe 可选）
        # 用 strip("|").split("|") 同时处理 "有/无尾部 pipe" 两种情况，
        # 避免 split("|")[1:-1] 在无尾部 pipe 时丢失最后一个单元格
        cells = [c.strip() for c in stripped.strip("|").split("|")]
        rows_text.append(cells)
        idx += 1

    if len(rows_text) < 1:
        # 完全没有表头行：把已消耗的行作为正文渲染
        for j in range(idx_init, idx):
            _add_body(doc, lines[j].strip())
        return idx

    # 即使只有表头（无数据行），也渲染表格（避免内容丢失）

    max_cols = max(len(r) for r in rows_text)
    table = doc.add_table(rows=len(rows_text), cols=max_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Light List Accent 1"
    table.autofit = True

    # 对齐方式默认值
    if col_aligns is None:
        col_aligns = [WD_ALIGN_PARAGRAPH.LEFT] * max_cols
    else:
        # 补齐到 max_cols
        while len(col_aligns) < max_cols:
            col_aligns.append(WD_ALIGN_PARAGRAPH.LEFT)

    for i, row_data in enumerate(rows_text):
        for j, cell_text in enumerate(row_data):
            cell = table.rows[i].cells[j]
            # 使用 _add_formatted_run 渲染（支持 **bold**、`code`、链接等）
            p = cell.paragraphs[0]
            _add_formatted_run(p, cell_text)
            for run in p.runs:
                run.font.size = Pt(9)
                run.font.name = "Microsoft YaHei"
            # 表头强制居中且加粗
            if i == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.bold = True
            else:
                # 数据行按 col_aligns 对齐
                p.alignment = col_aligns[j] if j < len(col_aligns) else WD_ALIGN_PARAGRAPH.LEFT
    return idx


# ── 主函数 ──────────────────────────────────────────────


def _render_cover(doc: Document, title: str, subtitle: str | None) -> None:
    """渲染封面：标题（必选）+ 副标题（可选）。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    p.paragraph_format.space_before = Pt(60)
    p.paragraph_format.space_after = Pt(20)

    if subtitle:
        p_sub = doc.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = p_sub.add_run(subtitle)
        sub_run.font.size = Pt(14)
        sub_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        sub_run.font.name = "Microsoft YaHei"
        sub_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        p_sub.paragraph_format.space_after = Pt(40)


def _has_heading(md_text: str) -> bool:
    """检测 markdown 文本中是否包含至少一个标题（# 开头且有内容的行）。"""
    for line in md_text.split("\n"):
        stripped = line.strip()
        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            if 1 <= level <= 9:
                # 去除 # 后必须还有内容才算标题
                content_after_hash = stripped.lstrip("#").strip()
                if content_after_hash:
                    return True
    return False


def _render_toc(doc: Document) -> None:
    """渲染目录页：标题 + TOC 字段 + 分页符。"""
    doc.add_paragraph("")  # 空行
    toc_heading = doc.add_heading("目录", level=1)
    for run in toc_heading.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    doc.element.body.append(parse_xml(_TOC_XML))
    doc.add_paragraph("").add_run().add_break(WD_BREAK.PAGE)


def _add_caption(doc: Document, text: str, *, color: tuple[int, int, int] = (0x66, 0x66, 0x66)) -> None:
    """添加居中 caption 段落（用于图片下方说明）。"""
    cap_p = doc.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap_p.add_run(text)
    cap_run.font.size = Pt(9)
    cap_run.font.color.rgb = RGBColor(*color)
    cap_run.font.name = "Microsoft YaHei"
    cap_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


@dataclass
class _DocState:
    """导出过程中的可变状态。"""
    chart_map: dict[int, Path]
    download_cache_dir: Path
    image_counter: int = 0
    chapter_count: int = 0


def _embed_image_with_caption(
    doc: Document, img_path: Path, alt_text: str, state: _DocState,
) -> None:
    """插入图片并附加 caption（仅成功插入时编号）。

    失败时清理已创建的空段落。
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    inserted = False
    try:
        pic_w, pic_h = _fit_image_size(img_path)
        run.add_picture(str(img_path), width=pic_w, height=pic_h)
        inserted = True
    except (OSError, ValueError, UnrecognizedImageError) as e:
        logger.warning("insert image failed: %s", e)
    if inserted:
        state.image_counter += 1
        cap_text = f"图 {state.image_counter}"
        if alt_text:
            cap_text += f": {alt_text}"
        _add_caption(doc, cap_text)
    else:
        # 插入失败：移除已创建的空段落，避免文档出现空白行
        try:
            p._element.getparent().remove(p._element)
        except (AttributeError, OSError):
            pass


def _embed_chart_with_caption(doc: Document, img_path: Path, title: str, state: _DocState) -> None:
    """插入 chart 图并附加 caption（含章节标题）。

    失败时清理已创建的空 caption 段落。
    """
    caption = None
    try:
        pic_w, pic_h = _fit_image_size(img_path)
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = caption.add_run()
        run.add_picture(str(img_path), width=pic_w, height=pic_h)
        state.image_counter += 1
        _add_caption(doc, f"图 {state.image_counter}: {title}")
    except (OSError, ValueError, UnrecognizedImageError) as e:
        logger.warning("  embed chart failed: %s", e)
        if caption is not None:
            try:
                caption._element.getparent().remove(caption._element)
            except (AttributeError, OSError):
                pass


def _process_markdown_line(
    doc: Document, lines: list[str], i: int, state: _DocState,
    *,
    skip_toc_heading: bool = False,
) -> int | None:
    """处理单行 markdown，返回下一个 idx；返回 None 表示行未被识别（按正文处理）。

    主入口 dispatch：图片、代码块、分隔线、引用、表格、标题、列表、任务列表。

    Args:
        skip_toc_heading: 如果已自动生成目录页，则跳过正文中 "# 目录" 标题，
                          避免目录标题重复。
    """
    stripped = lines[i].strip()

    if not stripped:
        return i + 1

    # 图片行 ![alt](path) 或 ![alt](http(s)://...)
    # URL 支持一层嵌套括号（与 _add_formatted_run 的 link 正则保持一致）
    m_img = re.match(
        r"^!\[([^\]]*)\]\(((?:[^()]|\([^()]*\))*)\)\s*$", stripped,
    )
    if m_img:
        alt_text = m_img.group(1)
        img_url_or_path = m_img.group(2)
        if img_url_or_path.startswith(("http://", "https://")):
            img_path = _download_image(img_url_or_path, state.download_cache_dir)
            if img_path is None:
                logger.warning("图片下载失败，作为正文显示: %s", img_url_or_path[:80])
                _add_body(doc, stripped)
                return i + 1
        else:
            img_path = Path(img_url_or_path)

        if img_path.exists():
            _embed_image_with_caption(doc, img_path, alt_text, state)
        else:
            _add_body(doc, stripped)
        return i + 1

    # 代码块（围栏 ```）
    if stripped.startswith("```"):
        return _parse_code_block(doc, lines, i)

    # 水平分隔线 ---/***/___ → 分页符
    if re.match(r"^(-{3,}|\*{3,}|_{3,})$", stripped):
        p = doc.add_paragraph()
        run = p.add_run()
        run.add_break(WD_BREAK.PAGE)
        return i + 1

    # 引用块（blockquote）：连续的 > 行合并为一段
    if stripped.startswith(">"):
        quote_lines: list[str] = []
        while i < len(lines) and lines[i].strip().startswith(">"):
            q_text = re.sub(r"^>\s*", "", lines[i].strip())
            quote_lines.append(q_text)
            i += 1
        merged = " ".join(quote_lines).strip()
        if merged:
            _add_blockquote(doc, merged)
        return i

    # markdown 表格：至少 2 行（表头 + 分隔行）才算表格
    # 表头尾部 | 可选（GFM 规范），靠下一行是否为分隔行判断
    if (
        stripped.startswith("|")
        and i + 1 < len(lines)
        and _parse_alignment_row(lines[i + 1].strip()) is not None
    ):
        return _parse_markdown_table(doc, stripped, lines, i)

    # 标题行（用 Word Heading 样式）
    if stripped.startswith("#"):
        level = min(len(stripped) - len(stripped.lstrip("#")), 9)
        # 去除首尾 # 与空白（支持 "## 标题 ##" 这种 atx 风格）
        text = stripped.lstrip("#").rstrip("#").strip()
        # 如果已自动生成目录页，跳过正文中的 "# 目录" 标题，避免重复
        if skip_toc_heading and level == 1 and text == "目录":
            return i + 1
        _add_heading(doc, text, level)
        i += 1
        if level <= 2:
            state.chapter_count += 1
            if state.chapter_count in state.chart_map:
                img_path = state.chart_map[state.chapter_count]
                if img_path.exists():
                    _embed_chart_with_caption(doc, img_path, text, state)
        return i

    # 列表项：支持嵌套（按行首空格计算层级）
    raw_line = lines[i]
    content_after_indent = raw_line.lstrip()
    indent_expanded = raw_line.expandtabs(4)
    leading_expanded = len(indent_expanded) - len(indent_expanded.lstrip())
    list_level = leading_expanded // 2

    # 任务列表项：- [ ] / - [x]
    m_task = re.match(r"^[-*]\s+\[([ xX])\]\s+(.*)$", content_after_indent)
    if m_task:
        checked = m_task.group(1) in ("x", "X")
        _add_task_list_item(doc, m_task.group(2), checked=checked, level=list_level)
        return i + 1

    # 无序列表项
    if content_after_indent.startswith("- ") or content_after_indent.startswith("* "):
        _add_bullet(doc, content_after_indent[2:], level=list_level)
        return i + 1

    # 有序列表：剥掉 "1. " 或 "1、" 前缀
    m = re.match(r"^(\d+)[.、]\s*(.*)$", content_after_indent)
    if m:
        _add_numbered(doc, m.group(2), level=list_level)
        return i + 1

    return None  # 未识别，交由正文处理


def _render_body(
    doc: Document, full_content: str, chart_images: list[tuple[int, Path]],
    download_cache_dir: Path,
    *,
    skip_toc_heading: bool = False,
) -> None:
    """渲染正文：逐行解析 markdown。"""
    chart_map: dict[int, Path] = {idx: path for idx, path in chart_images}
    state = _DocState(chart_map=chart_map, download_cache_dir=download_cache_dir)

    lines = full_content.split("\n")
    i = 0
    while i < len(lines):
        next_i = _process_markdown_line(doc, lines, i, state,
                                         skip_toc_heading=skip_toc_heading)
        if next_i is not None:
            i = next_i
            continue
        # 普通正文（支持 **bold**）
        _add_body(doc, lines[i].strip())
        i += 1


def _set_core_properties(
    doc: Document, title: str, author: str | None, subject: str | None,
) -> None:
    """设置文档元数据。"""
    cp = doc.core_properties
    cp.title = title
    cp.subject = subject or title
    cp.author = author or "Hermes AI"
    cp.comments = f"Hermes AI Report - 由 export_docx 生成于 {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    cp.created = datetime.now()
    cp.modified = datetime.now()


def _save_docx_atomic(doc: Document, output_path: Path) -> None:
    """原子保存 docx：先保存到 .tmp 再 rename，避免被 Word 占用导致半成品文件。"""
    tmp_path = output_path.with_suffix(output_path.suffix + ".tmp")
    try:
        doc.save(str(tmp_path))
        tmp_path.replace(output_path)
    except PermissionError as e:
        if tmp_path.exists():
            tmp_path.unlink()
        raise PermissionError(
            f"无法写入 {output_path}：文件可能被 Word 或其他程序占用。"
            f"请关闭后重试。"
        ) from e


def export_to_docx(
    title: str,
    full_content: str,
    chart_images: list[tuple[int, Path]],
    output_path: Path,
    *,
    subtitle: str | None = None,
    toc: bool = True,
    author: str | None = None,
    subject: str | None = None,
) -> Path:
    """将报告内容导出为 .docx 文件。

    使用 Word 内置 Heading 样式，支持自动生成目录。

    Args:
        title: 报告标题
        full_content: 报告全文（markdown 格式）
        chart_images: [(chapter_index, image_path), ...] chapter_index 从 1 开始
        output_path: 输出 .docx 文件路径
        subtitle: 封面副标题（可选）
        toc: 是否插入目录页，默认 True
        author: 文档作者（写入 core_properties，默认 "Hermes AI"）
        subject: 文档主题/描述（写入 core_properties，默认取 title）

    Returns:
        输出文件路径
    """
    doc = Document()

    # 默认字体
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Microsoft YaHei"
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    # ── 封面 ──
    _render_cover(doc, title, subtitle)

    # ── 目录（可选） ──
    # 只有 toc=True 且文档包含标题时才生成目录页
    # 避免无标题文档生成空目录导致"目录重复"
    toc_rendered = False
    if toc and _has_heading(full_content):
        _render_toc(doc)
        toc_rendered = True

    # ── 正文 ──
    download_cache_dir = output_path.parent / ".downloaded_images"
    _render_body(doc, full_content, chart_images, download_cache_dir,
                 skip_toc_heading=toc_rendered)

    # ── 文档元数据 ──
    _set_core_properties(doc, title, author, subject)

    # ── 原子写入 ──
    _save_docx_atomic(doc, output_path)
    logger.info("docx exported: %s", output_path)
    return output_path
