#!/usr/bin/env python3
"""数据中台一期实施计划 -> DOCX (含配图)"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── 路径 ──
MD_PATH = Path("/mnt/c/Users/1/Desktop/AI/报告团队/ai_report_system/reports/数据中台一期实施计划/数据中台一期实施计划.md")
CHARTS_DIR = Path("/mnt/c/Users/1/Desktop/AI/报告团队/ai_report_system/reports/数据中台一期实施计划/charts")
OUTPUT = Path("/mnt/c/Users/1/Desktop/AI/材料/数据中台一期实施计划_v2.docx")

# 配图映射：section_title_prefix -> chart_file
CHART_MAP = {
    "二、组织架构": CHARTS_DIR / "组织架构.png",
    "1.3 数据仓库": CHARTS_DIR / "数仓架构.png",
    "三、实施路线图": CHARTS_DIR / "路线图.png",
    "四、工作项依赖": CHARTS_DIR / "依赖关系.png",
}

with open(MD_PATH, "r", encoding="utf-8") as f:
    text = f.read()

doc = Document()

# ── 设置默认字体 ──
style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

# 设置标题字体
for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Microsoft YaHei'
    hs.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

# 页面设置
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── 封面 ──
for _ in range(6):
    doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run("数据中台一期实施计划")
run.font.size = Pt(26)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
run.font.name = 'Microsoft YaHei'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

subtitle_p = doc.add_paragraph()
subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle_p.add_run("IDM芯片制造企业 · 科研生产数据链贯通项目")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run.font.name = 'Microsoft YaHei'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

doc.add_paragraph()
info_p = doc.add_paragraph()
info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info_p.add_run("一期范围：订单履行域 | 总工期：19个月\n规划投资：约1,361万元（中台系统部分）")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.font.name = 'Microsoft YaHei'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

doc.add_page_break()

# ── 辅助函数 ──
def add_run(text, bold=False, size=Pt(10.5), color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = size
    r.font.bold = bold
    r.font.name = 'Microsoft YaHei'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    if color:
        r.font.color.rgb = color
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    return p

def add_table(headers, rows):
    """创建格式化表格"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ''
        p = hdr_cells[i].paragraphs[0]
        r = p.add_run(h)
        r.font.bold = True
        r.font.size = Pt(9)
        r.font.name = 'Microsoft YaHei'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 灰色背景
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E2F3"/>')
        hdr_cells[i]._tc.get_or_add_tcPr().append(shading)

    # 数据行
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = ''
            p = row_cells[i].paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(9)
            r.font.name = 'Microsoft YaHei'
            r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    return table


# ── 解析Markdown并构建DOCX ──
lines = text.split('\n')
i = 0
heading_history = []  # 跟踪标题层级，用于在合适位置插入图片

def get_current_heading_text():
    """获取当前H2/H3标题文本"""
    for h_type, h_text, h_level in reversed(heading_history):
        if h_level <= 2:
            return h_text
    return ""

while i < len(lines):
    line = lines[i]
    stripped = line.strip()

    # ── 处理标题 ──
    if stripped.startswith('#') and not stripped.startswith('# '):
        # markdown heading
        level = len(stripped.split()[0])
        h_text = stripped.lstrip('#').strip()
        heading_history.append((stripped, h_text, level))

        if level == 1:
            doc.add_heading(h_text, level=1)
        elif level == 2:
            doc.add_heading(h_text, level=2)
        elif level == 3:
            doc.add_heading(h_text, level=3)
        elif level >= 4:
            doc.add_heading(h_text, level=4)

        # ── 检查是否需要插入配图 ──
        for key, chart_path in CHART_MAP.items():
            if key in h_text and chart_path.exists():
                doc.add_paragraph()  # 空行
                img_p = doc.add_paragraph()
                img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = img_p.add_run()
                run.add_picture(str(chart_path), width=Inches(5.5))
                # 图注
                caption_p = doc.add_paragraph()
                caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = caption_p.add_run(f"图：{h_text}示意图")
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                r.font.name = 'Microsoft YaHei'
                r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                doc.add_paragraph()

        i += 1
        continue

    # ── 处理分割线 ──
    if re.match(r'^-{3,}$', stripped):
        doc.add_paragraph()  # 空行
        doc.add_paragraph('─' * 30)
        doc.add_paragraph()
        i += 1
        continue

    # ── 处理引用块 ──
    if stripped.startswith('>'):
        text_content = stripped.lstrip('>').strip()
        p = doc.add_paragraph()
        r = p.add_run(text_content)
        r.font.size = Pt(10)
        r.font.italic = True
        r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        r.font.name = 'Microsoft YaHei'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        # 灰色左边框
        pPr = p._p.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            '  <w:left w:val="single" w:sz="6" w:space="8" w:color="CCCCCC"/>'
            '</w:pBdr>'
        )
        pPr.append(pBdr)
        i += 1
        continue

    # ── 处理代码块（文本依赖关系图） ──
    if stripped.startswith('```'):
        i += 1
        code_lines = []
        while i < len(lines) and not lines[i].strip().startswith('```'):
            code_lines.append(lines[i])
            i += 1
        if code_lines:
            code_text = '\n'.join(code_lines)
            p = doc.add_paragraph()
            r = p.add_run(code_text)
            r.font.size = Pt(8)
            r.font.name = 'Courier New'
            # 灰色背景
            pPr = p._p.get_or_add_pPr()
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:fill="F5F5F5"/>')
            pPr.append(shading)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
        i += 1
        continue

    # ── 处理表格 ──
    if stripped.startswith('|') and (i + 1 < len(lines) and lines[i+1].strip().startswith('|') and '---' in lines[i+1].strip()):
        # 收集表格行
        table_rows = []
        while i < len(lines) and lines[i].strip().startswith('|'):
            table_rows.append(lines[i].strip())
            i += 1
        if len(table_rows) >= 2:
            # 第0行=表头，第1行=分隔符，2+行=数据
            headers = [c.strip().strip('|').strip() for c in table_rows[0].split('|') if c.strip()]
            data_rows = []
            for row_text in table_rows[2:]:
                cells = [c.strip().strip('|').strip() for c in row_text.split('|')]
                # 跳过空行
                if cells and any(c for c in cells):
                    data_rows.append(cells)
            if headers and data_rows:
                try:
                    add_table(headers, data_rows)
                    doc.add_paragraph()
                except Exception:
                    pass
        continue

    # ── 普通段落 ──
    if stripped:
        # 处理**粗体**标记
        parts = re.split(r'(\*\*.*?\*\*)', stripped)
        p = doc.add_paragraph()
        for part in parts:
            if part.startswith('**') and part.endswith('**') and len(part) > 4:
                inner = part[2:-2]
                r = p.add_run(inner)
                r.font.bold = True
            else:
                r = p.add_run(part)
            r.font.size = Pt(10.5)
            r.font.name = 'Microsoft YaHei'
            r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        p.paragraph_format.line_spacing = Pt(18)
    else:
        # 空行保持
        pass

    i += 1


# ── 保存 ──
doc.save(str(OUTPUT))
print(f"✅ DOCX exported: {OUTPUT}")
print(f"   大小: {OUTPUT.stat().st_size / 1024:.0f} KB")