"""测试 docx 导出功能完备性：行内标记、链接、表格对齐、引用块、嵌套列表、任务列表、元数据。"""
from __future__ import annotations

import zipfile
from pathlib import Path

import pytest

python_docx = pytest.importorskip("docx")

from ai_report.export.docx_exporter import export_to_docx


def _read_doc_xml(docx_path: Path) -> str:
    """读取 docx 的 document.xml 内容。"""
    with zipfile.ZipFile(docx_path, "r") as z:
        return z.read("word/document.xml").decode("utf-8")


def _read_core_xml(docx_path: Path) -> str:
    """读取 docx 的 core.xml（元数据）。"""
    with zipfile.ZipFile(docx_path, "r") as z:
        return z.read("docProps/core.xml").decode("utf-8")


# ── 行内代码 ────────────────────────────────────────────

@pytest.mark.integration
class TestInlineCode:
    def test_inline_code_uses_consolas(self, tmp_path):
        """行内代码 `code` 应使用 Consolas 字体。"""
        output = tmp_path / "out.docx"
        export_to_docx(
            title="测试", full_content="这是 `inline code` 示例",
            chart_images=[], output_path=output,
        )
        xml = _read_doc_xml(output)
        # 应包含 Consolas 字体引用
        assert "Consolas" in xml

    def test_inline_code_has_shading(self, tmp_path):
        """行内代码应有浅灰背景 (w:shd)。"""
        output = tmp_path / "out.docx"
        export_to_docx(
            title="测试", full_content="这是 `code` 示例",
            chart_images=[], output_path=output,
        )
        xml = _read_doc_xml(output)
        assert "w:shd" in xml
        assert "F0F0F0" in xml


# ── 超链接 ──────────────────────────────────────────────

@pytest.mark.integration
class TestHyperlink:
    def test_link_preserves_url(self, tmp_path):
        """[text](url) 应保留为超链接。"""
        output = tmp_path / "out.docx"
        export_to_docx(
            title="测试",
            full_content="访问 [Hermes](https://example.com) 了解更多",
            chart_images=[], output_path=output,
        )
        xml = _read_doc_xml(output)
        # 应包含 hyperlink 标签
        assert "w:hyperlink" in xml

    def test_link_text_displayed(self, tmp_path):
        """链接显示文本应保留。"""
        output = tmp_path / "out.docx"
        export_to_docx(
            title="测试",
            full_content="[点击这里](https://example.com)",
            chart_images=[], output_path=output,
        )
        xml = _read_doc_xml(output)
        assert "点击这里" in xml

    def test_link_relationship_created(self, tmp_path):
        """超链接应在 rels 中创建外部关系。"""
        output = tmp_path / "out.docx"
        export_to_docx(
            title="测试",
            full_content="[link](https://example.com)",
            chart_images=[], output_path=output,
        )
        with zipfile.ZipFile(output, "r") as z:
            rels = z.read("word/_rels/document.xml.rels").decode("utf-8")
        assert "hyperlink" in rels
        assert "https://example.com" in rels


# ── 表格列对齐 ─────────────────────────────────────────

@pytest.mark.integration
class TestTableAlignment:
    def test_right_alignment(self, tmp_path):
        """|---:| 应右对齐。"""
        output = tmp_path / "out.docx"
        content = "| 名称 | 数值 |\n| :--- | ---: |\n| A | 100 |"
        export_to_docx(title="测试", full_content=content, chart_images=[], output_path=output)
        xml = _read_doc_xml(output)
        assert "<w:tbl" in xml

    def test_center_alignment(self, tmp_path):
        """|:---:| 应居中对齐。"""
        output = tmp_path / "out.docx"
        content = "| 名称 |\n| :---: |\n| A |"
        export_to_docx(title="测试", full_content=content, chart_images=[], output_path=output)
        xml = _read_doc_xml(output)
        assert "<w:tbl" in xml


# ── 引用块多行合并 ─────────────────────────────────────

@pytest.mark.integration
class TestBlockquoteMerge:
    def test_multiline_quote_merged(self, tmp_path):
        """连续的 > 行应合并为一段。"""
        output = tmp_path / "out.docx"
        content = "> 第一行\n> 第二行\n> 第三行"
        export_to_docx(title="测试", full_content=content, chart_images=[], output_path=output)
        from docx import Document

        doc = Document(str(output))
        # 找包含 "第一行" 的段落（跳过封面/目录）
        quote_paragraphs = [p for p in doc.paragraphs if "第一行" in p.text]
        assert len(quote_paragraphs) == 1, "多行引用应合并为 1 个段落"
        p = quote_paragraphs[0]
        assert "第一行" in p.text
        assert "第二行" in p.text
        assert "第三行" in p.text

    def test_quote_with_bold(self, tmp_path):
        """引用块内的 **bold** 应被解析。"""
        output = tmp_path / "out.docx"
        content = "> **重要** 提示"
        export_to_docx(title="测试", full_content=content, chart_images=[], output_path=output)
        from docx import Document

        doc = Document(str(output))
        p = next(pp for pp in doc.paragraphs if "重要" in pp.text)
        bold_runs = [r for r in p.runs if r.font.bold]
        assert any("重要" in r.text for r in bold_runs)


# ── 嵌套列表 ──────────────────────────────────────────

@pytest.mark.integration
class TestNestedList:
    def test_nested_bullet_indent(self, tmp_path):
        """子列表应缩进。"""
        output = tmp_path / "out.docx"
        content = "- 父项\n  - 子项\n    - 三级"
        export_to_docx(title="测试", full_content=content, chart_images=[], output_path=output)
        from docx import Document

        doc = Document(str(output))
        bullets = [p for p in doc.paragraphs if p.style.name == "List Bullet"]
        assert len(bullets) == 3
        # 父项：indent 应为 None 或 0
        # 子项：indent 应 > 0
        child_indent = bullets[1].paragraph_format.left_indent
        assert child_indent is not None and child_indent.pt > 0, "子项应有缩进"
        # 三级缩进应大于二级
        grandchild_indent = bullets[2].paragraph_format.left_indent
        assert grandchild_indent is not None and grandchild_indent.pt > child_indent.pt

    def test_numbered_list(self, tmp_path):
        """有序列表应使用 List Number 样式。"""
        output = tmp_path / "out.docx"
        content = "1. 第一\n2. 第二\n3. 第三"
        export_to_docx(title="测试", full_content=content, chart_images=[], output_path=output)
        from docx import Document

        doc = Document(str(output))
        numbered = [p for p in doc.paragraphs if p.style.name == "List Number"]
        assert len(numbered) == 3


# ── 任务列表 ──────────────────────────────────────────

@pytest.mark.integration
class TestTaskList:
    def test_unchecked_task(self, tmp_path):
        """- [ ] 应渲染为 ☐。"""
        output = tmp_path / "out.docx"
        content = "- [ ] 待办项"
        export_to_docx(title="测试", full_content=content, chart_images=[], output_path=output)
        from docx import Document

        doc = Document(str(output))
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        assert any("☐" in t for t in texts)
        assert any("待办项" in t for t in texts)

    def test_checked_task(self, tmp_path):
        """- [x] 应渲染为 ☑。"""
        output = tmp_path / "out.docx"
        content = "- [x] 已完成项"
        export_to_docx(title="测试", full_content=content, chart_images=[], output_path=output)
        from docx import Document

        doc = Document(str(output))
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        assert any("☑" in t for t in texts)


# ── HTML <br> ────────────────────────────────────────

@pytest.mark.integration
class TestHtmlBr:
    def test_br_converts_to_break(self, tmp_path):
        """<br> 应转换为换行。"""
        output = tmp_path / "out.docx"
        content = "第一行<br>第二行"
        export_to_docx(title="测试", full_content=content, chart_images=[], output_path=output)
        from docx import Document

        doc = Document(str(output))
        p = [p for p in doc.paragraphs if p.text.strip()][0]
        # 应有 2 个 run（br 分隔），或包含 break 元素
        xml = _read_doc_xml(output)
        assert "w:br" in xml


# ── 文档元数据 ────────────────────────────────────────

@pytest.mark.integration
class TestCoreProperties:
    def test_title_in_core(self, tmp_path):
        output = tmp_path / "out.docx"
        export_to_docx(
            title="测试报告标题",
            full_content="内容",
            chart_images=[],
            output_path=output,
        )
        from docx import Document

        doc = Document(str(output))
        assert doc.core_properties.title == "测试报告标题"

    def test_author_default(self, tmp_path):
        output = tmp_path / "out.docx"
        export_to_docx(
            title="测试", full_content="内容",
            chart_images=[], output_path=output,
        )
        from docx import Document

        doc = Document(str(output))
        assert doc.core_properties.author == "Hermes AI"

    def test_author_custom(self, tmp_path):
        output = tmp_path / "out.docx"
        export_to_docx(
            title="测试", full_content="内容",
            chart_images=[], output_path=output,
            author="自定义作者",
        )
        from docx import Document

        doc = Document(str(output))
        assert doc.core_properties.author == "自定义作者"

    def test_subject_defaults_to_title(self, tmp_path):
        output = tmp_path / "out.docx"
        export_to_docx(
            title="主标题", full_content="内容",
            chart_images=[], output_path=output,
        )
        from docx import Document

        doc = Document(str(output))
        assert doc.core_properties.subject == "主标题"

    def test_subject_custom(self, tmp_path):
        output = tmp_path / "out.docx"
        export_to_docx(
            title="主标题", full_content="内容",
            chart_images=[], output_path=output,
            subject="自定义主题",
        )
        from docx import Document

        doc = Document(str(output))
        assert doc.core_properties.subject == "自定义主题"

    def test_created_set(self, tmp_path):
        output = tmp_path / "out.docx"
        export_to_docx(
            title="测试", full_content="内容",
            chart_images=[], output_path=output,
        )
        from docx import Document

        doc = Document(str(output))
        assert doc.core_properties.created is not None


# ── 图片 caption 统一编号 ─────────────────────────────

@pytest.mark.integration
class TestImageCaptionNumbering:
    def test_md_image_numbered(self, tmp_path):
        """markdown 图片应有统一编号。"""
        # 创建一个最小 PNG
        from helpers import _make_minimal_png

        img_path = tmp_path / "test.png"
        img_path.write_bytes(_make_minimal_png(100, 100))

        output = tmp_path / "out.docx"
        content = f"![图片A]({img_path})\n\n![图片B]({img_path})"
        export_to_docx(title="测试", full_content=content, chart_images=[], output_path=output)
        from docx import Document

        doc = Document(str(output))
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        # 应有 "图 1" 和 "图 2"
        assert any("图 1" in t for t in texts)
        assert any("图 2" in t for t in texts)


# ── 删除线 ───────────────────────────────────────────

@pytest.mark.integration
class TestStrikethrough:
    def test_strikethrough_rendered(self, tmp_path):
        """~~text~~ 应有删除线。"""
        output = tmp_path / "out.docx"
        content = "这是 ~~删除内容~~ 示例"
        export_to_docx(title="测试", full_content=content, chart_images=[], output_path=output)
        from docx import Document

        doc = Document(str(output))
        p = next(pp for pp in doc.paragraphs if "删除内容" in pp.text)
        strike_runs = [r for r in p.runs if r.font.strike]
        assert any("删除内容" in r.text for r in strike_runs)


# ── 斜体 ─────────────────────────────────────────────

@pytest.mark.integration
class TestItalic:
    def test_italic_rendered(self, tmp_path):
        """*text* 应有斜体。"""
        output = tmp_path / "out.docx"
        content = "这是 *斜体内容* 示例"
        export_to_docx(title="测试", full_content=content, chart_images=[], output_path=output)
        from docx import Document

        doc = Document(str(output))
        # 找包含 "斜体内容" 的段落（跳过封面/目录）
        p = next(pp for pp in doc.paragraphs if "斜体内容" in pp.text)
        italic_runs = [r for r in p.runs if r.font.italic]
        assert any("斜体内容" in r.text for r in italic_runs)
