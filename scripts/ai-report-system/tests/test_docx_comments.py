"""测试 docx_comments 模块。

构造带批注的 docx 文件需要直接操作 OOXML，python-docx 没有直接 API。
这里使用最小化构造 + 实际文件测试。
"""
from __future__ import annotations

import io
import sys
import zipfile
from pathlib import Path

import pytest

python_docx = pytest.importorskip("docx")
lxml_etree = pytest.importorskip("lxml.etree")

TESTS_DIR = Path(__file__).resolve().parent
PROJECT_DIR = TESTS_DIR.parent
sys.path.insert(0, str(PROJECT_DIR / "src"))
sys.path.insert(0, str(PROJECT_DIR / "scripts"))

from docx_comments import extract_chapter_comments, extract_comments


def _build_docx_with_comments(
    paragraphs: list[tuple[str, str | None, list[str]]],
) -> bytes:
    """构造一个带批注的 docx 文件。

    paragraphs: [(text, heading_level_or_None, [comment_text, ...]), ...]
    heading_level: 1 = H1, None = 普通段落
    """
    from docx import Document
    from lxml import etree as lxml_etree

    doc = Document()
    for text, level, comments in paragraphs:
        if level == 1:
            p = doc.add_heading(text, level=1)
        else:
            p = doc.add_paragraph(text)
        # 在段落上注入批注引用（简化实现）
        for cid, _ in enumerate(comments):
            # 添加 commentReference
            run = p.add_run()
            ref_xml = (
                f'<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                f'<w:rPr><w:rStyle w:val="CommentReference"/></w:rPr>'
                f'<w:commentReference w:id="{cid}"/></w:r>'
            )
            run._element.append(lxml_etree.fromstring(ref_xml))

    # 序列化到内存
    buf = io.BytesIO()
    doc.save(buf)

    # 重新打开 zip，注入 comments.xml
    buf.seek(0)
    original = zipfile.ZipFile(buf, "r")
    items = {n: original.read(n) for n in original.namelist()}
    original.close()

    # 构造 comments.xml
    all_comments = []
    for _, _, comments in paragraphs:
        all_comments.extend(comments)

    comments_xml_parts = ['<?xml version="1.0" encoding="UTF-8" standalone="yes"?>']
    comments_xml_parts.append(
        '<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
    )
    for cid, ctext in enumerate(all_comments):
        comments_xml_parts.append(
            f'<w:comment w:id="{cid}" w:author="test" w:date="2026-01-01T00:00:00Z" w:initials="t">'
            f'<w:p><w:r><w:t>{ctext}</w:t></w:r></w:p>'
            f'</w:comment>'
        )
    comments_xml_parts.append('</w:comments>')
    items["word/comments.xml"] = "".join(comments_xml_parts).encode("utf-8")

    # 添加关系到 .rels
    rels_path = "word/_rels/document.xml.rels"
    rels_xml = items[rels_path].decode("utf-8")
    if "comments.xml" not in rels_xml:
        new_rel = (
            '<Relationship Id="rIdComments" '
            'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments" '
            'Target="comments.xml"/>'
        )
        rels_xml = rels_xml.replace("</Relationships>", new_rel + "</Relationships>")
        items[rels_path] = rels_xml.encode("utf-8")

    # 添加 content-types
    ct_path = "[Content_Types].xml"
    ct_xml = items[ct_path].decode("utf-8")
    if "comments.xml" not in ct_xml:
        new_ct = (
            '<Override PartName="/word/comments.xml" '
            'ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"/>'
        )
        ct_xml = ct_xml.replace("</Types>", new_ct + "</Types>")
        items[ct_path] = ct_xml.encode("utf-8")

    # 重新打包
    out_buf = io.BytesIO()
    with zipfile.ZipFile(out_buf, "w", zipfile.ZIP_DEFLATED) as z:
        for name, data in items.items():
            z.writestr(name, data)
    return out_buf.getvalue()


class TestExtractComments:
    def test_no_comments_returns_empty(self, tmp_path):
        """无批注的 docx 返回空列表。"""
        from docx import Document

        docx_path = tmp_path / "empty.docx"
        Document().save(str(docx_path))
        assert extract_comments(docx_path) == []

    def test_extract_comments(self, tmp_path):
        docx_bytes = _build_docx_with_comments([
            ("第一章", 1, ["这是批注1"]),
            ("正文内容", None, []),
        ])
        docx_path = tmp_path / "test.docx"
        docx_path.write_bytes(docx_bytes)

        comments = extract_comments(docx_path)
        assert len(comments) == 1
        assert "批注1" in comments[0]["text"]


class TestExtractChapterComments:
    def test_empty_docx(self, tmp_path):
        from docx import Document

        docx_path = tmp_path / "empty.docx"
        Document().save(str(docx_path))
        assert extract_chapter_comments(docx_path) == []

    def test_comments_grouped_by_chapter(self, tmp_path):
        docx_bytes = _build_docx_with_comments([
            ("第一章", 1, ["批注A", "批注B"]),
            ("正文", None, []),
            ("第二章", 1, ["批注C"]),
        ])
        docx_path = tmp_path / "test.docx"
        docx_path.write_bytes(docx_bytes)

        result = extract_chapter_comments(docx_path)
        assert len(result) == 2
        # 第一章有 2 条批注
        assert result[0]["chapter_title"] == "第一章"
        assert result[0]["comment_count"] == 2
        # 第二章有 1 条批注
        assert result[1]["chapter_title"] == "第二章"
        assert result[1]["comment_count"] == 1

    def test_comments_preserve_order(self, tmp_path):
        """批注顺序应保留文档出现顺序，不按字母排序。"""
        docx_bytes = _build_docx_with_comments([
            ("章节", 1, ["zebra", "apple", "mango"]),
        ])
        docx_path = tmp_path / "test.docx"
        docx_path.write_bytes(docx_bytes)

        result = extract_chapter_comments(docx_path)
        assert len(result) == 1
        comments = result[0]["comments"]
        # 不应按字母排序
        assert comments[0] == "zebra"
        assert comments[1] == "apple"
        assert comments[2] == "mango"

    def test_include_full_content(self, tmp_path):
        docx_bytes = _build_docx_with_comments([
            ("第一章", 1, ["批注"]),
            ("正文内容XYZ", None, []),
        ])
        docx_path = tmp_path / "test.docx"
        docx_path.write_bytes(docx_bytes)

        result = extract_chapter_comments(docx_path, include_content=True)
        assert len(result) == 1
        assert "正文内容XYZ" in result[0]["full_content"]

    def test_dedup_same_comment_text(self, tmp_path):
        """同章节重复批注文本应去重。"""
        docx_bytes = _build_docx_with_comments([
            ("章节", 1, ["重复批注", "重复批注"]),
        ])
        docx_path = tmp_path / "test.docx"
        docx_path.write_bytes(docx_bytes)

        result = extract_chapter_comments(docx_path)
        assert len(result) == 1
        assert result[0]["comment_count"] == 1
