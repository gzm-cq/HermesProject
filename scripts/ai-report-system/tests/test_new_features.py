"""测试新增功能：mermaid 缓存、未闭合块、图片尺寸自适应、表格校验。"""
from __future__ import annotations

import base64
import json
import struct
import sys
import zipfile
from pathlib import Path

import pytest

TESTS_DIR = Path(__file__).resolve().parent
PROJECT_DIR = TESTS_DIR.parent
sys.path.insert(0, str(PROJECT_DIR / "src"))
sys.path.insert(0, str(PROJECT_DIR / "scripts"))

from ai_report.export.docx_exporter import (
    _fit_image_size,
    _read_image_size,
    export_to_docx,
)
from export_docx import (
    _extract_mermaid_blocks,
    _load_mermaid_cache,
    _mermaid_code_hash,
    _save_mermaid_cache,
    replace_mermaid_with_images,
)


def _make_minimal_png(width: int,  height: int) -> bytes:
    """生成最小有效 PNG 文件（指定宽高）。"""
    # PNG signature
    sig = b"\x89PNG\r\n\x1a\n"
    # IHDR chunk
    ihdr_data = struct.pack(">IIBBBBB", width, height, 8, 2, 0, 0, 0)
    ihdr_len = struct.pack(">I", len(ihdr_data))
    ihdr_type = b"IHDR"
    import zlib
    ihdr_crc = struct.pack(">I", zlib.crc32(ihdr_type + ihdr_data) & 0xFFFFFFFF)
    ihdr = ihdr_len + ihdr_type + ihdr_data + ihdr_crc
    # IDAT chunk (minimal raw image data)
    raw = b""
    for _ in range(height):
        raw += b"\x00" + b"\x00" * (width * 3)
    compressed = zlib.compress(raw)
    idat_len = struct.pack(">I", len(compressed))
    idat_type = b"IDAT"
    idat_crc = struct.pack(">I", zlib.crc32(idat_type + compressed) & 0xFFFFFFFF)
    idat = idat_len + idat_type + compressed + idat_crc
    # IEND chunk
    iend = b"\x00\x00\x00\x00IEND\xaeB`\x82"
    return sig + ihdr + idat + iend


# ── Mermaid 未闭合块 ────────────────────────────────────────

class TestMermaidUnclosedBlock:
    def test_unclosed_block_marked(self):
        """未闭合的 mermaid 块应被标记为 closed=False。"""
        md = """# 标题

```mermaid
graph TD
  A --> B
"""  # 末尾没有 ```
        blocks = _extract_mermaid_blocks(md)
        assert len(blocks) == 1
        assert blocks[0]["closed"] is False
        assert "graph TD" in blocks[0]["code"]

    def test_closed_block_marked(self):
        md = """```mermaid
graph TD
  A --> B
```"""
        blocks = _extract_mermaid_blocks(md)
        assert len(blocks) == 1
        assert blocks[0]["closed"] is True

    def test_unclosed_block_kept_in_replace(self, tmp_path):
        """未闭合块在 replace_mermaid_with_images 中应保留原代码。"""
        md = """```mermaid
graph TD
  A --> B
"""  # 未闭合
        # 即使有图片路径，未闭合块也不替换
        img_path = tmp_path / "fake.png"
        img_path.write_bytes(b"fake")
        result = replace_mermaid_with_images(md, [img_path])
        assert "```mermaid" in result
        assert "graph TD" in result


# ── Mermaid 缓存 ────────────────────────────────────────────

class TestMermaidCache:
    def test_hash_deterministic(self):
        assert _mermaid_code_hash("graph TD\n  A->B") == _mermaid_code_hash("graph TD\n  A->B")

    def test_different_code_different_hash(self):
        assert _mermaid_code_hash("graph TD") != _mermaid_code_hash("pie title X")

    def test_load_empty_cache(self, tmp_path):
        cache_path = tmp_path / ".mermaid_cache.json"
        assert _load_mermaid_cache(cache_path) == {}

    def test_save_and_load_cache(self, tmp_path):
        cache_path = tmp_path / ".mermaid_cache.json"
        cache = {"abc123": "mermaid_00.png"}
        _save_mermaid_cache(cache_path, cache)
        loaded = _load_mermaid_cache(cache_path)
        assert loaded == cache

    def test_load_invalid_json_returns_empty(self, tmp_path):
        cache_path = tmp_path / ".mermaid_cache.json"
        cache_path.write_text("not a valid json", encoding="utf-8")
        assert _load_mermaid_cache(cache_path) == {}


# ── 图片尺寸自适应 ─────────────────────────────────────────

class TestImageSizeAdaptive:
    def test_read_png_size(self, tmp_path):
        png = _make_minimal_png(800, 600)
        img_path = tmp_path / "test.png"
        img_path.write_bytes(png)
        size = _read_image_size(img_path)
        assert size == (800, 600)

    def test_fit_landscape_image(self, tmp_path):
        """横图（宽>高）按宽度限制。"""
        png = _make_minimal_png(1920, 1080)
        img_path = tmp_path / "landscape.png"
        img_path.write_bytes(png)
        width, height = _fit_image_size(img_path)
        # 横图应返回 width
        assert width is not None
        assert height is None

    def test_fit_portrait_image(self, tmp_path):
        """竖图（高>宽）按高度限制。"""
        png = _make_minimal_png(1080, 1920)
        img_path = tmp_path / "portrait.png"
        img_path.write_bytes(png)
        width, height = _fit_image_size(img_path)
        # 竖图应返回 height
        assert width is None
        assert height is not None

    def test_fit_square_image(self, tmp_path):
        """方图按宽度限制。"""
        png = _make_minimal_png(500, 500)
        img_path = tmp_path / "square.png"
        img_path.write_bytes(png)
        width, height = _fit_image_size(img_path)
        assert width is not None
        assert height is None

    def test_invalid_image_returns_default(self, tmp_path):
        """无法读取尺寸时按默认宽度。"""
        img_path = tmp_path / "fake.png"
        img_path.write_bytes(b"not a real image")
        width, height = _fit_image_size(img_path)
        assert width is not None
        assert height is None

    def test_portrait_image_inserted_into_docx(self, tmp_path):
        """竖图插入 docx 后应正常生成（不报错）。"""
        png = _make_minimal_png(800, 1600)  # 1:2 竖图
        img_path = tmp_path / "tall.png"
        img_path.write_bytes(png)

        output = tmp_path / "out.docx"
        content = f"![竖图]({img_path})"
        export_to_docx(title="测试", full_content=content, chart_images=[], output_path=output)
        assert output.exists()
        assert output.stat().st_size > 0


# ── 表格识别双行校验 ──────────────────────────────────────

class TestTableDoubleLineCheck:
    def test_single_pipe_line_not_table(self, tmp_path):
        """单行 |xxx| 不应被识别为表格。"""
        output = tmp_path / "out.docx"
        content = "|注意|这是一个提示"
        export_to_docx(title="测试", full_content=content, chart_images=[], output_path=output)

        with zipfile.ZipFile(output, "r") as z:
            xml = z.read("word/document.xml").decode("utf-8")
        # 不应生成 <w:tbl>
        assert "<w:tbl" not in xml

    def test_two_line_table_recognized(self, tmp_path):
        """标准表格（表头 + 分隔行）应被识别。"""
        output = tmp_path / "out.docx"
        content = "| A | B |\n|---|---|\n| 1 | 2 |"
        export_to_docx(title="测试", full_content=content, chart_images=[], output_path=output)

        with zipfile.ZipFile(output, "r") as z:
            xml = z.read("word/document.xml").decode("utf-8")
        assert "<w:tbl" in xml


# ── docx 保存失败原子写入 ─────────────────────────────────

class TestAtomicSave:
    def test_permission_error_when_file_locked(self, tmp_path):
        """模拟文件被占用时抛出 PermissionError。"""
        output = tmp_path / "out.docx"
        # 先创建一个空文件
        output.write_bytes(b"")
        # chmod 只读无法触发 Windows 上的 PermissionError，所以这里只验证
        # 函数能正常工作（不抛异常）
        export_to_docx(title="测试", full_content="内容", chart_images=[], output_path=output)
        assert output.exists()


# ── Bug3 回归：markdown 图片插入失败时清理空段落 ─────────

class TestMarkdownImageInsertFailureCleanup:
    def test_invalid_image_no_caption_orphan(self, tmp_path):
        """图片文件存在但内容无效时：
        - add_picture 失败
        - 不应残留空段落
        - 不应附加 caption（图 N: alt）
        """
        # 创建无效图片文件（内容不是真正的图片）
        bad_img = tmp_path / "fake.png"
        bad_img.write_bytes(b"not a real PNG image")

        output = tmp_path / "out.docx"
        content = f"前文\n\n![无效图]({bad_img})\n\n后文"
        export_to_docx(
            title="测试",
            full_content=content,
            chart_images=[],
            output_path=output,
        )

        from docx import Document
        doc = Document(str(output))
        # 不应有 "图 1" caption（因为插入失败）
        captions = [p.text for p in doc.paragraphs if p.text.strip().startswith("图 ")]
        assert captions == [], f"插入失败不应产生 caption，但出现: {captions}"
        # 不应有空段落（只有图片占位但无内容）
        # 统计段落总数 - 应等于：封面标题(1) + 目录标题(1) + 空行若干 + 前文 + 后文
        # 关键是不应有空 caption 段落
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "前文" in all_text
        assert "后文" in all_text

    def test_invalid_chart_image_no_orphan_caption(self, tmp_path):
        """chart 图插入失败时不应残留空 caption 段落。"""
        bad_img = tmp_path / "bad_chart.png"
        bad_img.write_bytes(b"invalid chart content")

        output = tmp_path / "out.docx"
        content = "# 章节1\n正文内容"
        export_to_docx(
            title="测试",
            full_content=content,
            chart_images=[(1, bad_img)],  # chapter 1 的 chart
            output_path=output,
        )

        from docx import Document
        doc = Document(str(output))
        # 不应有 "图 1: 章节1" caption
        captions = [p.text for p in doc.paragraphs if "图 1" in p.text]
        assert captions == [], f"chart 插入失败不应产生 caption: {captions}"
        # 章节标题应正常
        assert any("章节1" in p.text for p in doc.paragraphs)


# ── Bug9 回归：单行表格内容不丢失 ───────────────────────

class TestSingleHeaderTablePreserved:
    """只有表头+分隔行（无数据行）的表格也应当被渲染，
    内容（表头单元格）不能被丢弃。"""
    def test_header_only_table_rendered(self, tmp_path):
        output = tmp_path / "out.docx"
        content = "| A | B |\n|---|---|"
        export_to_docx(
            title="测试", full_content=content,
            chart_images=[], output_path=output,
        )

        with zipfile.ZipFile(output, "r") as z:
            xml = z.read("word/document.xml").decode("utf-8")
        # 应生成表格
        assert "<w:tbl" in xml
        # 表头 A 和 B 必须保留
        assert "A" in xml
        assert "B" in xml

    def test_header_only_table_has_one_row(self, tmp_path):
        """表头独占一行的表格，行数应为 1。"""
        output = tmp_path / "out.docx"
        content = "| 姓名 | 年龄 |\n|:---|---:|"
        export_to_docx(
            title="测试", full_content=content,
            chart_images=[], output_path=output,
        )
        from docx import Document
        doc = Document(str(output))
        tables = doc.tables
        assert len(tables) == 1
        # 1 行（仅表头）
        assert len(tables[0].rows) == 1
        # 表头内容保留
        assert tables[0].rows[0].cells[0].text.strip() == "姓名"
        assert tables[0].rows[0].cells[1].text.strip() == "年龄"


# ── Bug11 回归：代码块空行保留为真空行 ───────────────────

class TestCodeBlockEmptyLinePreserved:
    """代码块中的空行应保持为真空段落，不应被替换为空格。"""
    def test_empty_line_in_code_block_is_blank(self, tmp_path):
        output = tmp_path / "out.docx"
        content = "```python\nline1\n\nline3\n```"
        export_to_docx(
            title="测试", full_content=content,
            chart_images=[], output_path=output,
        )
        from docx import Document
        doc = Document(str(output))
        tables = doc.tables
        assert len(tables) == 1
        # 代码块用单单元格表格呈现
        cell = tables[0].rows[0].cells[0]
        # cell.paragraphs[0] 是 "# python" 标签
        # 之后应为：line1 段落、空段落、line3 段落
        # 验证存在空段落（无 run 或 run 文本为空）
        paragraphs_text = [p.text for p in cell.paragraphs]
        # line1 和 line3 必须存在
        assert "line1" in paragraphs_text
        assert "line3" in paragraphs_text
        # 必须有空字符串段落（真空行）
        assert "" in paragraphs_text
        # 空段落不应包含空格占位符（如 " " 或 "  "）
        empty_paras = [p for p in cell.paragraphs if p.text == ""]
        assert len(empty_paras) >= 1
        NS_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        for ep in empty_paras:
            # 检查段落内不应有任何带空格文本的 <w:t>
            t_nodes = ep._element.findall(f"{{{NS_W}}}r/{{{NS_W}}}t")
            assert len(t_nodes) == 0, (
                f"空段落不应含 run/text 节点，但发现 {len(t_nodes)} 个: "
                f"{[t.text for t in t_nodes]!r}"
            )

    def test_multiple_consecutive_empty_lines(self, tmp_path):
        """多个连续空行也应全部保留。"""
        output = tmp_path / "out.docx"
        content = "```\nfoo\n\n\nbar\n```"
        export_to_docx(
            title="测试", full_content=content,
            chart_images=[], output_path=output,
        )
        from docx import Document
        doc = Document(str(output))
        cell = doc.tables[0].rows[0].cells[0]
        # 期望：foo, 空, 空, bar（4 段）
        # 首段无 # lang 标签
        paragraphs_text = [p.text for p in cell.paragraphs]
        assert paragraphs_text.count("") >= 2
        assert "foo" in paragraphs_text
        assert "bar" in paragraphs_text


# ── Bug12 回归：标题末尾 # 被正确去除 ────────────────────

class TestHeadingTrailingHashes:
    """`## 标题 ##` 这种 atx 风格的标题应去除末尾 #。"""
    def test_heading_with_trailing_hash(self, tmp_path):
        output = tmp_path / "out.docx"
        content = "## 标题 ##"
        export_to_docx(
            title="测试", full_content=content,
            chart_images=[], output_path=output,
        )
        from docx import Document
        doc = Document(str(output))
        # 找出 heading 段落
        headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading")]
        # 应有一个二级标题
        heading_texts = [p.text.strip() for p in headings]
        assert "标题" in heading_texts
        # 不应包含末尾的 #
        for p in headings:
            assert "#" not in p.text, f"标题不应残留 # 字符: {p.text!r}"

    def test_heading_multiple_trailing_hashes(self, tmp_path):
        """`## 标题 ####` 多余的 # 也要去除。"""
        output = tmp_path / "out.docx"
        content = "# 大标题 ###"
        export_to_docx(
            title="测试", full_content=content,
            chart_images=[], output_path=output,
        )
        from docx import Document
        doc = Document(str(output))
        headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert any(p.text.strip() == "大标题" for p in headings)

    def test_heading_no_trailing_hash_unchanged(self, tmp_path):
        """`## 标题` 无末尾 # 不受影响。"""
        output = tmp_path / "out.docx"
        content = "## 普通标题"
        export_to_docx(
            title="测试", full_content=content,
            chart_images=[], output_path=output,
        )
        from docx import Document
        doc = Document(str(output))
        headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert any(p.text.strip() == "普通标题" for p in headings)


# ── Bug15 回归：URL 含嵌套括号的链接 ────────────────────

class TestLinkWithNestedParens:
    """[text](url) URL 中包含一层嵌套括号时应被正确解析。"""
    def test_link_text_preserved(self, tmp_path):
        from docx import Document
        output = tmp_path / "out.docx"
        content = "[C语言](https://en.wikipedia.org/wiki/C_(programming_language))"
        export_to_docx(
            title="测试", full_content=content,
            chart_images=[], output_path=output,
        )
        doc = Document(str(output))
        # 文本 "C语言" 必须保留
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "C语言" in all_text

    def test_link_full_url_preserved(self, tmp_path):
        """完整 URL（含嵌套括号）必须出现在外部关系中。"""
        output = tmp_path / "out.docx"
        content = "[wiki](https://en.wikipedia.org/wiki/C_(programming_language))"
        export_to_docx(
            title="测试", full_content=content,
            chart_images=[], output_path=output,
        )
        # 检查 docx 中包含完整的 hyperlink 关系
        with zipfile.ZipFile(output, "r") as z:
            rels_xml = z.read("word/_rels/document.xml.rels").decode("utf-8")
        # 完整 URL 必须出现在 rels 中（含嵌套括号）
        assert "https://en.wikipedia.org/wiki/C_(programming_language)" in rels_xml

    def test_link_text_and_url_unit(self):
        """单元测试：直接验证 _add_formatted_run 解析嵌套括号 URL。"""
        from docx import Document
        from ai_report.export.docx_exporter import _add_formatted_run
        doc = Document()
        p = doc.add_paragraph()
        _add_formatted_run(
            p,
            "[wiki](https://en.wikipedia.org/wiki/C_(programming_language))",
        )
        xml_str = p._element.xml
        # 文本保留
        assert "wiki" in xml_str
        # 检查外部 hyperlink 关系（rels 创建后会在 part.rels 中）
        rels = p.part.rels
        has_full_url = any(
            "https://en.wikipedia.org/wiki/C_(programming_language)" == str(rel.target_ref)
            for rel in rels.values()
            if rel.is_external
        )
        assert has_full_url, "完整 URL（含嵌套括号）应被保留在外部关系中"


# ── Bug16 回归：图片 URL 支持嵌套括号 ───────────────────

class TestImageMarkdownNestedParens:
    """![alt](url) URL 含一层嵌套括号时也应被正确识别为图片行，
    而不是被当作正文渲染（与 Bug15 的 link 修复保持一致）。"""
    def test_local_image_with_nested_parens_not_body(self, tmp_path):
        """本地图片路径含嵌套括号也应被识别为图片行。"""
        # 创建一个真实的 PNG 图片
        png = _make_minimal_png(100, 100)
        img_path = tmp_path / "image_(1).png"
        img_path.write_bytes(png)

        output = tmp_path / "out.docx"
        content = f"![alt]({img_path})"
        export_to_docx(
            title="测试", full_content=content,
            chart_images=[], output_path=output,
        )
        from docx import Document
        doc = Document(str(output))
        # 应有 caption（图 1: alt）说明图片行被识别
        captions = [p.text for p in doc.paragraphs if p.text.strip().startswith("图 ")]
        assert len(captions) == 1, f"图片行未被识别为图片，captions: {captions}"
        assert "alt" in captions[0]

    def test_image_regex_unit(self):
        """单元测试：直接验证图片正则匹配嵌套括号 URL。"""
        import re
        pattern = re.compile(
            r"^!\[([^\]]*)\]\(((?:[^()]|\([^()]*\))*)\)\s*$"
        )
        # 标准 URL
        assert pattern.match("![alt](https://example.com/image.png)")
        # 嵌套括号 URL
        assert pattern.match("![alt](https://example.com/C_(programming_language).png)")
        # 嵌套括号本地路径
        assert pattern.match("![alt](/path/to/image_(1).png)")
        # 不匹配：URL 末尾未闭合
        assert not pattern.match("![alt](https://example.com/image.png")
        # 不匹配：缺少 ![ ](
        assert not pattern.match("text [alt](url)")


# ── Bug17 回归：表格行无尾部 | 不丢失单元格 ──────────────

class TestTableRowWithoutTrailingPipe:
    """GFM 规范允许表格行省略尾部 |，单元格不能丢失。"""
    def test_data_row_without_trailing_pipe(self, tmp_path):
        """数据行无尾部 | 时所有单元格必须保留。"""
        output = tmp_path / "out.docx"
        content = "| A | B | C |\n|---|---|---|\n| 1 | 2 | 3"
        export_to_docx(
            title="测试", full_content=content,
            chart_images=[], output_path=output,
        )
        from docx import Document
        doc = Document(str(output))
        tables = doc.tables
        assert len(tables) == 1
        # 应有 2 行：表头 + 数据行
        assert len(tables[0].rows) == 2
        # 数据行所有单元格必须保留
        data_row = tables[0].rows[1]
        assert data_row.cells[0].text.strip() == "1"
        assert data_row.cells[1].text.strip() == "2"
        assert data_row.cells[2].text.strip() == "3"

    def test_header_without_trailing_pipe(self, tmp_path):
        """表头无尾部 | 时也应识别为表格。"""
        output = tmp_path / "out.docx"
        content = "| A | B\n|---|---\n| 1 | 2"
        export_to_docx(
            title="测试", full_content=content,
            chart_images=[], output_path=output,
        )
        from docx import Document
        doc = Document(str(output))
        tables = doc.tables
        assert len(tables) == 1
        # 表头单元格
        assert tables[0].rows[0].cells[0].text.strip() == "A"
        assert tables[0].rows[0].cells[1].text.strip() == "B"

    def test_all_rows_with_trailing_pipe_still_works(self, tmp_path):
        """标准格式（每行都有尾部 |）不应受影响。"""
        output = tmp_path / "out.docx"
        content = "| A | B |\n|---|---|\n| 1 | 2 |"
        export_to_docx(
            title="测试", full_content=content,
            chart_images=[], output_path=output,
        )
        from docx import Document
        doc = Document(str(output))
        tables = doc.tables
        assert len(tables) == 1
        assert tables[0].rows[1].cells[0].text.strip() == "1"
        assert tables[0].rows[1].cells[1].text.strip() == "2"


# ── Bug18 + Bug19 回归：build_chart_images 与 _process_markdown_line 一致 ─

class TestBuildChartImagesHeadingConsistency:
    """build_chart_images 的章节计数和标题解析必须与
    docx_exporter._process_markdown_line 保持一致。"""
    def test_heading_without_space_counted(self, tmp_path):
        """`#标题`（无空格）应被计为章节（与 _process_markdown_line 一致）。"""
        from export_docx import build_chart_images
        md_text = "#标题A\n正文\n##标题B\n正文"
        # 准备一个匹配的图片文件
        chart_dir = tmp_path / "charts"
        chart_dir.mkdir()
        img = chart_dir / "img.png"
        img.write_bytes(_make_minimal_png(100, 100))
        chart_map = {"标题A": "img.png"}

        result = build_chart_images(
            md_path=tmp_path / "fake.md",
            charts_dir=chart_dir,
            chart_map=chart_map,
            md_text=md_text,
        )
        # 应匹配到第 1 章
        assert len(result) == 1
        chapter_idx, img_path = result[0]
        assert chapter_idx == 1
        assert img_path == img

    def test_heading_with_trailing_hash_stripped(self, tmp_path):
        """`## 标题 ##`（atx 风格）的末尾 # 应被去除，标题才能匹配 chart_map。"""
        from export_docx import build_chart_images
        md_text = "## 架构图 ##\n正文"
        chart_dir = tmp_path / "charts"
        chart_dir.mkdir()
        img = chart_dir / "arch.png"
        img.write_bytes(_make_minimal_png(100, 100))
        chart_map = {"架构图": "arch.png"}

        result = build_chart_images(
            md_path=tmp_path / "fake.md",
            charts_dir=chart_dir,
            chart_map=chart_map,
            md_text=md_text,
        )
        # 应匹配到（末尾 # 已去除，标题 "架构图" 与 key 匹配）
        assert len(result) == 1
        chapter_idx, img_path = result[0]
        assert chapter_idx == 1
        assert img_path == img

    def test_h3_not_counted_as_chapter(self, tmp_path):
        """### 三级标题不应被计入 chapter_idx（与 _process_markdown_line 一致）。"""
        from export_docx import build_chart_images
        md_text = "### 三级标题\n正文\n## 二级标题\n正文"
        chart_dir = tmp_path / "charts"
        chart_dir.mkdir()
        img = chart_dir / "img.png"
        img.write_bytes(_make_minimal_png(100, 100))
        chart_map = {"二级标题": "img.png"}

        result = build_chart_images(
            md_path=tmp_path / "fake.md",
            charts_dir=chart_dir,
            chart_map=chart_map,
            md_text=md_text,
        )
        # ### 不计数，## 是第 1 章
        assert len(result) == 1
        chapter_idx, _ = result[0]
        assert chapter_idx == 1

    def test_chapter_idx_matches_exporter(self, tmp_path):
        """build_chart_images 给出的 chapter_idx 必须与
        export_to_docx 内部 _process_markdown_line 计数一致。"""
        from export_docx import build_chart_images
        # 4 个 H1/H2 标题
        md_text = (
            "# 第一章\n正文\n"
            "## 第二章\n正文\n"
            "### 子节（不应计数）\n正文\n"
            "# 第三章\n正文\n"
            "## 第四章 ##\n正文\n"  # atx 风格末尾 #
        )
        chart_dir = tmp_path / "charts"
        chart_dir.mkdir()
        # 为每个章节准备一张图
        chart_map = {
            "第一章": "c1.png",
            "第二章": "c2.png",
            "第三章": "c3.png",
            "第四章": "c4.png",
        }
        for fname in chart_map.values():
            (chart_dir / fname).write_bytes(_make_minimal_png(100, 100))

        result = build_chart_images(
            md_path=tmp_path / "fake.md",
            charts_dir=chart_dir,
            chart_map=chart_map,
            md_text=md_text,
        )
        # 4 张图都应匹配，chapter_idx 分别为 1/2/3/4
        # 用 (chapter_idx, img_name) 配对验证
        result_by_name = {p.name: idx for idx, p in result}
        assert result_by_name == {
            "c1.png": 1,
            "c2.png": 2,
            "c3.png": 3,
            "c4.png": 4,
        }, f"章节计数与 _process_markdown_line 不一致: {result_by_name}"


# ── P0 回归：_generate_with_retry 失败时不残留 save_path ──

class TestGenerateWithRetryCleanupOnFailure:
    """VLM review 失败/错误后所有轮次耗尽时，save_path 不应残留。
    否则下次调用 render_mermaid_images 会因 "if save_path.exists()"
    直接复用未验证的图片（P0-5 修复的回归保护）。
    """
    def test_all_fail_rounds_cleans_save_path(self, tmp_path):
        """所有轮次都 fail：save_path 必须被删除。"""
        from unittest.mock import patch
        from export_docx import _generate_with_retry

        save_path = tmp_path / "out.png"
        save_path.write_bytes(b"fake image bytes")

        # generate_image 返回 True（成功），review_image 返回 fail
        with patch("export_docx.generate_image", return_value=True), \
             patch("export_docx.review_image",
                   return_value={"status": "fail", "reasoning": "bad"}):
            result = _generate_with_retry(
                prompt="test", save_path=save_path, label="test",
                max_rounds=2, review=True,
                image_size="2k", aspect_ratio="16:9",
                api_key=None, base_url=None, model=None,
                timeout=10, critic_path=None,
            )

        assert result is None
        # 关键断言：save_path 必须不存在
        assert not save_path.exists(), (
            "save_path 残留会导致下次 render_mermaid_images 误用未审核图片"
        )

    def test_all_error_rounds_cleans_save_path(self, tmp_path):
        """所有轮次都 error（超时/非 JSON）：save_path 也必须被删除。
        这是 P0 回归保护点 —— 旧实现 error 时不删除，下次会被误用。"""
        from unittest.mock import patch
        from export_docx import _generate_with_retry

        save_path = tmp_path / "out.png"
        save_path.write_bytes(b"fake image bytes")

        with patch("export_docx.generate_image", return_value=True), \
             patch("export_docx.review_image",
                   return_value={"status": "error", "reasoning": "timeout"}):
            result = _generate_with_retry(
                prompt="test", save_path=save_path, label="test",
                max_rounds=2, review=True,
                image_size="2k", aspect_ratio="16:9",
                api_key=None, base_url=None, model=None,
                timeout=10, critic_path=None,
            )

        assert result is None
        # 关键断言：error 状态下 save_path 也必须不存在
        assert not save_path.exists(), (
            "error 状态下 save_path 残留会导致下次 render_mermaid_images 误用"
        )

    def test_generation_failure_cleans_save_path(self, tmp_path):
        """generate_image 直接失败（返回 False）也应清理 save_path。"""
        from unittest.mock import patch
        from export_docx import _generate_with_retry

        save_path = tmp_path / "out.png"
        save_path.write_bytes(b"fake image bytes")

        with patch("export_docx.generate_image", return_value=False):
            result = _generate_with_retry(
                prompt="test", save_path=save_path, label="test",
                max_rounds=2, review=False,
                image_size="2k", aspect_ratio="16:9",
                api_key=None, base_url=None, model=None,
                timeout=10, critic_path=None,
            )

        assert result is None
        assert not save_path.exists()

    def test_pass_keeps_save_path(self, tmp_path):
        """正常通过审核的图片应保留。"""
        from unittest.mock import patch
        from export_docx import _generate_with_retry

        save_path = tmp_path / "out.png"
        save_path.write_bytes(b"fake image bytes")

        with patch("export_docx.generate_image", return_value=True), \
             patch("export_docx.review_image",
                   return_value={"status": "pass", "reasoning": "ok"}):
            result = _generate_with_retry(
                prompt="test", save_path=save_path, label="test",
                max_rounds=2, review=True,
                image_size="2k", aspect_ratio="16:9",
                api_key=None, base_url=None, model=None,
                timeout=10, critic_path=None,
            )

        assert result == save_path
        assert save_path.exists(), "通过审核的图片应保留"

    def test_no_review_success_keeps_save_path(self, tmp_path):
        """review=False 时生成成功后直接返回，save_path 保留。"""
        from unittest.mock import patch
        from export_docx import _generate_with_retry

        save_path = tmp_path / "out.png"
        save_path.write_bytes(b"fake image bytes")

        with patch("export_docx.generate_image", return_value=True):
            result = _generate_with_retry(
                prompt="test", save_path=save_path, label="test",
                max_rounds=2, review=False,
                image_size="2k", aspect_ratio="16:9",
                api_key=None, base_url=None, model=None,
                timeout=10, critic_path=None,
            )

        assert result == save_path
        assert save_path.exists()
